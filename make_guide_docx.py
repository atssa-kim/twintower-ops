"""
make_guide_docx.py
대원설정방법.md -> 대원설정방법.docx 변환
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── 페이지 여백 설정 ──────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = section.right_margin = Cm(2.5)
section.top_margin  = section.bottom_margin = Cm(2.0)

# ── 색상 ──────────────────────────────────────────────────────
RED    = RGBColor(0xA3, 0x2D, 0x2D)
BLUE   = RGBColor(0x0C, 0x44, 0x7C)
GRAY   = RGBColor(0x47, 0x55, 0x69)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LGRAY  = RGBColor(0xF1, 0xF5, 0xF9)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def para_style(p, size=11, bold=False, color=None, align=None, space_before=0, space_after=6):
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if align: pf.alignment = align
    for run in p.runs:
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        if color: run.font.color.rgb = color

def add_heading(text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 0:   # 제목
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = WHITE
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.CENTER
        # 배경 빨간 박스
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'A32D2D')
        pPr.append(shd)
        p.paragraph_format.left_indent  = Cm(-2.5)
        p.paragraph_format.right_indent = Cm(-2.5)
        pPr2 = p._p.get_or_add_pPr()
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '-1440')
        ind.set(qn('w:right'), '-1440')
        pPr2.append(ind)
    elif level == 1:  # ## 섹션
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = WHITE
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(4)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '0C447C')
        pPr.append(shd)
    elif level == 2:  # ### 소제목
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = BLUE
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(3)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    return p

def add_body(text, bullet=False, indent=0, bold=False, color=None, size=11):
    p = doc.add_paragraph()
    if bullet:
        p.paragraph_format.left_indent   = Cm(0.5 + indent * 0.5)
        p.paragraph_format.first_line_indent = Cm(-0.4)
        run0 = p.add_run('• ')
        run0.font.size  = Pt(size)
        run0.font.name  = '맑은 고딕'
        run0.font.color.rgb = RED
    else:
        p.paragraph_format.left_indent = Cm(indent * 0.5)
    run = p.add_run(text)
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.name  = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    if color: run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    return p

def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name  = 'Courier New'
    run.font.size  = Pt(10)
    run.font.color.rgb = BLUE
    p.paragraph_format.left_indent  = Cm(1.0)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), 'EEF2F7')
    pPr.append(shd)
    return p

def add_note(text):
    p = doc.add_paragraph()
    run = p.add_run('⚠️  ' + text)
    run.font.size  = Pt(10)
    run.font.name  = '맑은 고딕'
    run.font.color.rgb = RGBColor(0xBA, 0x75, 0x17)
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 헤더
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_bg(cell, '0C447C')
        run = cell.paragraphs[0].add_run(h)
        run.font.bold  = True
        run.font.size  = Pt(10)
        run.font.color.rgb = WHITE
        run.font.name  = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 데이터
    for r, row in enumerate(rows, 1):
        for c, val in enumerate(row):
            cell = table.cell(r, c)
            if r % 2 == 0:
                set_cell_bg(cell, 'F1F5F9')
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = '맑은 고딕'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    # 열 너비
    if col_widths:
        for r in range(len(rows) + 1):
            for c, w in enumerate(col_widths):
                table.cell(r, c).width = Cm(w)
    return table

def add_spacer(pt=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pt)


# ════════════════════════════════════════════════════
#  문서 본문
# ════════════════════════════════════════════════════

# 제목
add_heading('🔔  트윈타워 상황대응 앱\n대원 설정 방법', level=0)
add_spacer(6)
p = add_body('재난 발령 시 즉시 알림을 받으려면 아래 순서대로 설정하세요.', color=GRAY)
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_spacer(4)

# 1단계
add_heading('1단계.  앱 열기', level=1)
add_body('스마트폰 Chrome 브라우저에서 아래 주소를 입력합니다.')
add_code('https://atssa-kim.github.io/twintower-ops/')

# 2단계
add_heading('2단계.  홈 화면에 추가 (바탕화면 바로가기 설치)', level=1)
add_body('앱을 홈 화면에 설치해야 백그라운드 알림이 정상 동작합니다.')

add_heading('Android (Chrome)', level=2)
add_body('브라우저 우측 상단 ⋮ 메뉴 탭', bullet=True)
add_body('"앱 설치" 또는 "홈 화면에 추가" 선택', bullet=True)
add_body('설치 완료 후 홈 화면 아이콘으로 실행', bullet=True)

add_heading('iPhone / iPad (Safari)', level=2)
add_note('iOS는 반드시 Safari로 열어야 합니다. Chrome 앱으로는 설치 불가.')
add_body('Safari에서 위 주소 접속', bullet=True)
add_body('하단 공유 버튼 (□↑) 탭', bullet=True)
add_body('"홈 화면에 추가" 선택', bullet=True)
add_body('이름 확인 후 추가 탭', bullet=True)
add_body('홈 화면의 아이콘으로 실행  ← 이후 항상 아이콘으로 열 것', bullet=True, bold=True)
add_note('iPhone은 iOS 16.4 이상이어야 알림이 지원됩니다.')

# 3단계
add_heading('3단계.  로그인', level=1)
add_body('앱 실행', bullet=True)
add_body('사번 선택 — 본인 사번을 드롭다운에서 선택', bullet=True)
add_body('로그인 버튼 탭', bullet=True)

# 4단계
add_heading('4단계.  발령 알림 켜기', level=1)
add_body('로그인 후 화면 상단에 아래 배너가 표시됩니다.')
add_code('🔔 발령 알림 받기 — 눌러서 켜기')
add_body('노란 배너를 탭', bullet=True)
add_body('"알림 허용" 팝업 → 허용 선택', bullet=True)
add_body('"✅ 알림이 켜졌습니다" 메시지 확인', bullet=True)
add_note('이 단계를 완료해야 재난 발령 시 푸시 알림이 옵니다.')

# 5단계
add_heading('5단계.  설정 완료 확인', level=1)
add_table(
    ['항목', '확인 방법'],
    [
        ('앱 설치', '홈 화면에 🔔 아이콘 존재'),
        ('로그인',  '상단에 내 이름·소속 표시'),
        ('알림 설정', '노란 배너가 사라짐'),
    ],
    col_widths=[4, 11]
)
add_spacer()

# 발령 시 사용
add_heading('발령 시 사용 방법', level=1)
add_body('푸시 알림 수신 → 알림 탭 → 앱 자동 실행', bullet=True)
add_body('또는 앱 아이콘 직접 실행', bullet=True)
add_body('화면에 재난 발령 내용 및 내 임무 목록 자동 표시', bullet=True)
add_body('임무 수행 시 항목을 탭하여 체크 ✓', bullet=True, bold=True)
add_body('지휘자 화면에서 임무 수행률 실시간 확인 가능', bullet=True)

# OS별 지원
add_heading('스마트폰 OS별 지원 현황', level=1)
add_table(
    ['환경', '지원 여부', '비고'],
    [
        ('Android Chrome', '✅ 지원', '홈 화면 설치 권장'),
        ('iPhone iOS 16.4 이상 (Safari)', '✅ 지원', 'Safari → 홈 화면 추가 필수'),
        ('iPhone iOS 16.3 이하', '❌ 미지원', '기기 업데이트 필요'),
        ('iOS Safari 일반 브라우저', '❌ 알림 불가', '홈 화면 설치 후 아이콘으로 실행'),
    ],
    col_widths=[5.5, 3, 6.5]
)
add_spacer()

# FAQ
add_heading('자주 묻는 질문', level=1)

add_body('Q. 알림 배너가 안 보여요.', bold=True)
add_body('Firebase 연결(LIVE 모드)일 때만 표시됩니다. 상단 LIVE 배지를 확인하세요.', indent=1, color=GRAY)

add_body('Q. iPhone에서 알림이 안 와요.', bold=True)
add_body('Safari로 접속 후 홈 화면에 추가했는지 확인하세요. iOS 16.4 미만은 웹 푸시를 지원하지 않습니다.', indent=1, color=GRAY)

add_body('Q. 알림이 차단됐다고 나와요.', bold=True)
add_body('스마트폰 설정 → 알림 → 브라우저 → 해당 사이트 알림을 허용으로 변경하세요.', indent=1, color=GRAY)

add_body('Q. 사번이 목록에 없어요.', bold=True)
add_body('관리자에게 사번 등록을 요청하세요.', indent=1, color=GRAY)

add_spacer(10)
p = add_body('문의: 트윈타워 방재센터', color=GRAY, size=9)
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── 저장 ──────────────────────────────────────────
OUT = '대원설정방법.docx'
doc.save(OUT)
print(f'[OK] {OUT} 생성 완료')
